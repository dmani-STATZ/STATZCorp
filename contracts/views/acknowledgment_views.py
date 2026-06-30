import base64
import io
import logging
import uuid

from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.utils import timezone
from docx import Document

from STATZWeb.decorators import conditional_login_required
from contracts.services.sharepoint_paths import resolve_contract_folder_path
from contracts.services.sharepoint_service import (
    SharePointError,
    _get_drive_item,
    convert_file_to_pdf_bytes,
    delete_file_by_id,
    download_file_bytes_by_id,
    normalize_folder_path,
    open_file_in_browser,
    send_pdf_bytes_to_folder,
    upload_bytes_to_folder,
)
from ..models import AcknowledgementLetter, AcknowledgmentLetterTemplate, Clin
from ..forms import AcknowledgementLetterForm
from suppliers.contact_categories import PRIMARY_CATEGORY_NAME
from suppliers.models import Contact
from users.user_settings import UserSettings

logger = logging.getLogger('django')

ACK_TEMPLATE_FOLDER = 'Statz-Public/data/V87/aFed-DOD/z-temp/templates'
ACK_TEMP_FOLDER = 'Statz-Public/data/V87/aFed-DOD/z-temp'
DOCX_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
)


def _acknowledgment_pdf_filename(letter) -> str:
    return "Purchase Order Acknowledgment Letter.pdf"


def _build_letter_substitutions(letter) -> dict:
    """Build the substitution dict from a saved AcknowledgementLetter instance."""

    def fmt_date(val):
        if not val:
            return ''
        try:
            return val.strftime('%B %d, %Y')
        except AttributeError:
            return str(val)

    salutation = (letter.salutation or '').strip()
    fname = (letter.addr_fname or '').strip()
    lname = (letter.addr_lname or '').strip()
    recipient_name = f"{salutation} {fname} {lname}".strip()

    city = (letter.city or '').strip()
    state = (letter.state or '').strip()
    zip_ = (letter.zip or '').strip()
    city_state_zip = f"{city}, {state} {zip_}".strip().strip(',')

    po = (letter.po or '').strip()
    po_ext = (letter.po_ext or '').strip()
    po_number = f"{po}{f' / {po_ext}' if po_ext else ''}"

    return {
        '{{LETTER_DATE}}': fmt_date(letter.letter_date),
        '{{RECIPIENT_NAME}}': recipient_name,
        '{{SUPPLIER_NAME}}': letter.supplier or '',
        '{{STREET_ADDRESS}}': letter.st_address or '',
        '{{CITY_STATE_ZIP}}': city_state_zip,
        '{{PO_NUMBER}}': po_number,
        '{{CONTRACT_NUMBER}}': letter.contract_num or '',
        '{{SUPPLIER_DUE_DATE}}': fmt_date(letter.supplier_due_date),
        '{{FAT_DUE_DATE}}': fmt_date(letter.fat_due_date),
        '{{PLT_DUE_DATE}}': fmt_date(letter.plt_due_date),
        '{{DPAS_PRIORITY}}': letter.dpas_priority or '',
        '{{STATZ_CONTACT}}': letter.statz_contact or '',
        '{{STATZ_TITLE}}': letter.statz_contact_title or '',
        '{{STATZ_PHONE}}': letter.statz_contact_phone or '',
        '{{STATZ_EMAIL}}': letter.statz_contact_email or '',
    }


def _apply_substitutions_to_doc(doc, substitutions: dict) -> None:
    """Apply placeholder substitutions to a python-docx Document in place."""

    def replace_in_paragraph(para, placeholder, value):
        if placeholder not in para.text:
            return
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))
                return
        full_text = ''.join(r.text for r in para.runs)
        if placeholder in full_text and para.runs:
            para.runs[0].text = full_text.replace(placeholder, str(value))
            for r in para.runs[1:]:
                r.text = ''

    for placeholder, value in substitutions.items():
        for para in doc.paragraphs:
            replace_in_paragraph(para, placeholder, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, placeholder, value)


def _download_template_docx(active_template):
    if not active_template.sharepoint_file_id:
        raise SharePointError('Active template has no SharePoint file ID.')
    template_bytes = download_file_bytes_by_id(active_template.sharepoint_file_id)
    return Document(io.BytesIO(template_bytes))


def _generate_letter_pdf_bytes(letter) -> bytes:
    """
    Fill the active Word template from a saved letter and return PDF bytes
    via SharePoint temp upload + Graph conversion.
    """
    active_template = AcknowledgmentLetterTemplate.get_active()
    if not active_template:
        raise ValueError(
            'No active letter template. Ask a staff member to upload one.'
        )
    if not active_template.sharepoint_file_id:
        raise ValueError(
            'Active template has no SharePoint file ID. Re-upload the template.'
        )

    substitutions = _build_letter_substitutions(letter)
    try:
        doc = _download_template_docx(active_template)
        _apply_substitutions_to_doc(doc, substitutions)
        buffer = io.BytesIO()
        doc.save(buffer)
        filled_bytes = buffer.getvalue()
    except SharePointError:
        raise
    except Exception as e:
        logger.error("Letter substitution failed: %s", e)
        raise ValueError('Failed to prepare document.') from e

    temp_file_id = None
    try:
        temp_filename = f"ack-{uuid.uuid4().hex[:8]}.docx"
        temp_result = upload_bytes_to_folder(
            ACK_TEMP_FOLDER,
            temp_filename,
            filled_bytes,
            content_type=DOCX_CONTENT_TYPE,
        )
        temp_file_id = temp_result['id']
        return convert_file_to_pdf_bytes(temp_file_id)
    except SharePointError:
        raise
    finally:
        if temp_file_id:
            try:
                delete_file_by_id(temp_file_id)
            except Exception:
                logger.warning(
                    "Letter PDF cleanup: could not delete temp file %s",
                    temp_file_id,
                )


def _get_or_create_acknowledgment_letter(clin, request):
    """Get or create an AcknowledgementLetter for a CLIN with prefill logic."""
    letter = AcknowledgementLetter.objects.filter(clin=clin).first()
    if not letter:
        letter = AcknowledgementLetter(clin=clin)

    if letter.pk is not None:
        return letter

    if clin.supplier:
        supplier = clin.supplier
        letter.supplier = supplier.name

        contact = Contact.objects.filter(
            supplier=supplier,
            categories__name=PRIMARY_CATEGORY_NAME,
        ).first()
        if contact:
            letter.salutation = contact.salutation
            if contact.name:
                names = contact.name.split(maxsplit=1)
                letter.addr_fname = names[0]
                letter.addr_lname = names[1] if len(names) > 1 else ''
        else:
            letter.salutation = None
            letter.addr_fname = None
            letter.addr_lname = None

        if supplier.physical_address:
            addr = supplier.physical_address
            letter.st_address = addr.address_line_1
            letter.city = addr.city
            letter.state = addr.state
            letter.zip = addr.zip
        else:
            letter.st_address = None
            letter.city = None
            letter.state = None
            letter.zip = None
    else:
        letter.supplier = None
        letter.salutation = None
        letter.addr_fname = None
        letter.addr_lname = None
        letter.st_address = None
        letter.city = None
        letter.state = None
        letter.zip = None

    letter.po = clin.po_number
    letter.po_ext = clin.po_num_ext
    letter.contract_num = clin.contract.contract_number if clin.contract else None

    p_clin = (
        Clin.objects
        .filter(contract=clin.contract, item_type='P', supplier_due_date__isnull=False)
        .order_by('supplier_due_date')
        .first()
    )
    letter.supplier_due_date = p_clin.supplier_due_date if p_clin else None

    fat_clin = (
        Clin.objects
        .filter(contract=clin.contract, item_type__in=['C', 'G'], supplier_due_date__isnull=False)
        .order_by('supplier_due_date')
        .first()
    )
    letter.fat_due_date = fat_clin.supplier_due_date if fat_clin else None

    plt_clin = (
        Clin.objects
        .filter(contract=clin.contract, item_type='L', supplier_due_date__isnull=False)
        .order_by('supplier_due_date')
        .first()
    )
    letter.plt_due_date = plt_clin.supplier_due_date if plt_clin else None

    letter.statz_contact = f"{request.user.first_name} {request.user.last_name}".strip()
    letter.statz_contact_email = request.user.email

    user_settings = UserSettings.get_multiple_settings(request.user, [
        'statz_contact_title',
        'statz_contact_phone',
    ])
    letter.statz_contact_title = user_settings.get('statz_contact_title', 'Contract Manager')
    letter.statz_contact_phone = user_settings.get('statz_contact_phone', '')

    if not letter.pk:
        letter.letter_date = timezone.now().date()

    letter.save()
    return letter


def _template_download_url(template):
    if not template or not template.sharepoint_file_id:
        return ''
    try:
        return open_file_in_browser(template.sharepoint_file_id)
    except SharePointError:
        return ''


@conditional_login_required
def get_acknowledgment_letter(request, clin_id):
    """Get or create an acknowledgment letter for a CLIN (modal JSON API)."""
    clin = get_object_or_404(Clin, id=clin_id)
    letter = _get_or_create_acknowledgment_letter(clin, request)
    form = AcknowledgementLetterForm(instance=letter)

    return JsonResponse({
        'success': True,
        'html': render(request, 'contracts/partials/acknowledgment_letter_form.html', {
            'form': form,
            'clin': clin,
            'letter': letter
        }).content.decode('utf-8')
    })


@conditional_login_required
def acknowledgment_letter_page(request, clin_id):
    """Dedicated full-page editor for the PO acknowledgment letter."""
    clin = get_object_or_404(
        Clin.objects.select_related('contract', 'supplier', 'supplier__physical_address'),
        id=clin_id,
    )
    letter = _get_or_create_acknowledgment_letter(clin, request)
    form = AcknowledgementLetterForm(instance=letter)
    active_template = AcknowledgmentLetterTemplate.get_active()

    return render(request, 'contracts/acknowledgment_letter_page.html', {
        'letter': letter,
        'form': form,
        'clin': clin,
        'contract': clin.contract,
        'active_template': active_template,
        'template_download_url': _template_download_url(active_template),
        'user_is_staff': request.user.is_staff,
    })


@conditional_login_required
def preview_acknowledgment_letter(request, letter_id):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required.'})

    letter = get_object_or_404(AcknowledgementLetter, pk=letter_id)

    try:
        pdf_bytes = _generate_letter_pdf_bytes(letter)
        pdf_b64 = base64.b64encode(pdf_bytes).decode('utf-8')
        return JsonResponse({'success': True, 'pdf_base64': pdf_b64})
    except SharePointError as e:
        return JsonResponse({'success': False, 'error': e.message})
    except ValueError as e:
        return JsonResponse({'success': False, 'error': str(e)})


@conditional_login_required
def send_acknowledgment_to_contract_folder(request, letter_id):
    """
    Generate the filled PDF (same flow as preview) and save it directly
    to the contract's SharePoint folder. Overwrites if already exists.
    """
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required.'})

    letter = get_object_or_404(
        AcknowledgementLetter.objects.select_related('clin__contract'),
        pk=letter_id,
    )
    contract = letter.clin.contract

    active_template = AcknowledgmentLetterTemplate.get_active()
    if not active_template or not active_template.sharepoint_file_id:
        return JsonResponse({
            'success': False,
            'error': 'No active letter template. Ask a staff member to upload one.'
        })

    try:
        pdf_bytes = _generate_letter_pdf_bytes(letter)
    except SharePointError as e:
        return JsonResponse({'success': False, 'error': e.message})
    except ValueError as e:
        return JsonResponse({'success': False, 'error': str(e)})

    try:
        resolution = resolve_contract_folder_path(contract)
        contract_folder = resolution['path'].rstrip('/')

        pdf_filename = _acknowledgment_pdf_filename(letter)

        send_pdf_bytes_to_folder(contract_folder, pdf_filename, pdf_bytes)

        letter.is_user_edited = True
        letter.save(update_fields=['is_user_edited'])

        return JsonResponse({
            'success': True,
            'message': f'Saved to contract folder as {pdf_filename}',
            'filename': pdf_filename,
        })
    except SharePointError as e:
        return JsonResponse({'success': False, 'error': e.message})


@conditional_login_required
def get_existing_acknowledgment_pdf(request, letter_id):
    """
    Check if a previously saved acknowledgment PDF exists in the
    contract's SharePoint folder. If found, return it as base64.
    Returns success: False (not an error) if no file exists yet.
    """
    if request.method != 'GET':
        return JsonResponse({'success': False, 'error': 'GET required.'})

    letter = get_object_or_404(
        AcknowledgementLetter.objects.select_related('clin__contract'),
        pk=letter_id,
    )
    contract = letter.clin.contract
    pdf_filename = _acknowledgment_pdf_filename(letter)

    try:
        resolution = resolve_contract_folder_path(contract)
        contract_folder = resolution['path'].rstrip('/')
        file_path = normalize_folder_path(f"{contract_folder}/{pdf_filename}")

        item = _get_drive_item(file_path)
        if item is None:
            return JsonResponse({'success': False, 'exists': False})

        file_id = item.get('id')
        if not file_id:
            return JsonResponse({'success': False, 'exists': False})

        pdf_bytes = download_file_bytes_by_id(file_id)
        pdf_b64 = base64.b64encode(pdf_bytes).decode('utf-8')
        return JsonResponse({
            'success': True,
            'exists': True,
            'pdf_base64': pdf_b64,
            'filename': pdf_filename,
        })

    except SharePointError as e:
        logger.warning("get_existing_acknowledgment_pdf: %s", e.message)
        return JsonResponse({'success': False, 'exists': False})


@conditional_login_required
def upload_acknowledgment_template(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required.'})
    if not request.user.is_staff:
        return JsonResponse({'success': False, 'error': 'Staff only.'}, status=403)

    uploaded_file = request.FILES.get('template_file')
    rev_number = request.POST.get('rev_number', '').strip()

    if not uploaded_file:
        return JsonResponse({'success': False, 'error': 'No file provided.'})
    if not uploaded_file.name.endswith('.docx'):
        return JsonResponse({'success': False, 'error': 'File must be a .docx'})
    if not rev_number:
        return JsonResponse({'success': False, 'error': 'Revision number required.'})

    try:
        file_bytes = b''.join(uploaded_file.chunks())
        result = upload_bytes_to_folder(
            ACK_TEMPLATE_FOLDER,
            uploaded_file.name,
            file_bytes,
            content_type=DOCX_CONTENT_TYPE,
        )
    except SharePointError as e:
        return JsonResponse({'success': False, 'error': e.message})

    template = AcknowledgmentLetterTemplate(
        rev_number=rev_number,
        uploaded_by=request.user,
        sharepoint_file_id=result['id'],
        sharepoint_file_name=result['name'],
    )
    template.save()
    template.activate()

    download_url = result.get('downloadUrl') or _template_download_url(template)

    return JsonResponse({
        'success': True,
        'rev_number': template.rev_number,
        'uploaded_by': template.uploaded_by.get_full_name() or
                       template.uploaded_by.username,
        'uploaded_at': template.uploaded_at.strftime('%B %d, %Y'),
        'download_url': download_url,
    })


@conditional_login_required
def update_acknowledgment_letter(request, letter_id):
    """Update an existing acknowledgment letter."""
    letter = get_object_or_404(AcknowledgementLetter, id=letter_id)

    if request.method == 'POST':
        form = AcknowledgementLetterForm(request.POST, instance=letter)
        if form.is_valid():
            letter = form.save()

            statz_contact_title = form.cleaned_data.get('statz_contact_title')
            statz_contact_phone = form.cleaned_data.get('statz_contact_phone')

            if statz_contact_title:
                UserSettings.save_setting(
                    user=request.user,
                    name='statz_contact_title',
                    value=statz_contact_title,
                    setting_type='string',
                    description='User contact title for acknowledgment letters'
                )

            if statz_contact_phone:
                UserSettings.save_setting(
                    user=request.user,
                    name='statz_contact_phone',
                    value=statz_contact_phone,
                    setting_type='string',
                    description='User contact phone for acknowledgment letters'
                )

            return JsonResponse({'success': True})
        return JsonResponse({
            'success': False,
            'errors': form.errors
        })

    return JsonResponse({'success': False, 'error': 'Invalid request method'})
